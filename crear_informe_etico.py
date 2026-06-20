"""
Informe Ético y de Gobernanza — Modelo de Churn Olist
Grupo 7 · Sprint 4
Ejecutar: python crear_informe_etico.py
Requiere: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── COLORES ──────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
BLUE   = RGBColor(0x2E, 0x86, 0xAB)
ORANGE = RGBColor(0xF1, 0x8F, 0x01)
GREEN  = RGBColor(0x1A, 0xB3, 0x74)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0xF0, 0xF4, 0xF8)

doc = Document()

# ─── MÁRGENES ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ─── ESTILOS BASE ─────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    # Línea inferior azul
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2E86AB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def info_box(doc, text, color_hex='E8F4FD'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic    = True
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(40)
r = p_title.add_run('INFORME ÉTICO Y DE GOBERNANZA')
r.bold      = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Modelo de Predicción de Churn — Dataset Olist')
r2.font.size = Pt(14)
r2.font.color.rgb = BLUE

doc.add_paragraph()

# Tabla de metadatos
meta = doc.add_table(rows=4, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
filas_meta = [
    ('Proyecto',  'Predicción de Churn — E-commerce Olist'),
    ('Grupo',     'Grupo 7 — Percy Fuentes'),
    ('Sprint',    'Sprint 4 — Modelo Final'),
    ('Fecha',     datetime.datetime.today().strftime('%d de %B de %Y')),
]
for i, (k, v) in enumerate(filas_meta):
    row = meta.rows[i]
    set_cell_bg(row.cells[0], '1E3A5F')
    c0 = row.cells[0].paragraphs[0]
    r_k = c0.add_run(k)
    r_k.bold = True
    r_k.font.color.rgb = WHITE
    r_k.font.size = Pt(11)
    c1 = row.cells[1].paragraphs[0]
    r_v = c1.add_run(v)
    r_v.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════
heading1(doc, '1. Introducción y Contexto')
body(doc,
    'El presente informe analiza las implicaciones éticas, de privacidad y de gobernanza '
    'del modelo de predicción de churn desarrollado sobre el dataset público de Olist, '
    'una plataforma de e-commerce brasileña. El modelo predice qué clientes tienen mayor '
    'probabilidad de abandonar la plataforma en el siguiente mes, utilizando su historial '
    'de comportamiento de compra.')
body(doc,
    'Un modelo predictivo de esta naturaleza puede tener impactos directos sobre personas '
    'reales —los clientes— al determinar quiénes reciben ofertas de retención, descuentos '
    'o atención prioritaria. Por ello, es fundamental evaluar su uso desde una perspectiva '
    'ética y de responsabilidad.')

info_box(doc,
    'Modelo final: Logistic Regression (LogReg) | AUC Live: 63% | Lift Decil 1: 2.24x | '
    '27 variables de comportamiento | Datos: Olist 2016-2018 (público, Kaggle)',
    'E8F4FD')

# ══════════════════════════════════════════════════════════
# 2. TRANSPARENCIA Y EXPLICABILIDAD
# ══════════════════════════════════════════════════════════
heading1(doc, '2. Transparencia y Explicabilidad')

heading2(doc, '2.1 Elección de modelo interpretable')
body(doc,
    'Se eligió Logistic Regression como modelo final, en lugar de modelos de mayor '
    'complejidad como Gradient Boosting (HGB), precisamente por su interpretabilidad '
    'nativa. Los coeficientes del modelo permiten explicar directamente la dirección '
    'e impacto de cada variable sobre la predicción.')

heading2(doc, '2.2 Herramientas de explicabilidad aplicadas')
bullet(doc, 'Coeficientes normalizados: importancia relativa de cada variable (max 11.57% para comprador_unico).', 'Coeficientes LogReg')
bullet(doc, 'SHAP LinearExplainer: valores SHAP por cliente, beeswarm y bar plots.', 'SHAP')
bullet(doc, 'Permutation Importance: validación cruzada de la importancia de features.', 'Permutation Importance')
bullet(doc, 'Gain/Lift Tables: cuantificación del poder de discriminación por deciles.', 'Gain & Lift')

heading2(doc, '2.3 Narrativa explicable para el negocio')
body(doc, 'El modelo puede comunicarse a equipos no técnicos de la siguiente manera:')
info_box(doc,
    'Un cliente tiene mayor riesgo de churn si: lleva mucho tiempo sin comprar (recencia alta), '
    'compra siempre en la misma categoría (comprador_unico alto), recibe pedidos tarde '
    '(lead_min alto) y usa muchas cuotas (installments_max alto). '
    'Por el contrario, compras frecuentes y buenas reseñas reducen el riesgo.',
    'E8F8F0')

# ══════════════════════════════════════════════════════════
# 3. ANÁLISIS DE SESGOS
# ══════════════════════════════════════════════════════════
heading1(doc, '3. Análisis de Sesgos Potenciales')

heading2(doc, '3.1 Sesgo de selección temporal')
body(doc,
    'El modelo fue entrenado con datos de 2016 a enero 2018, período de rápido crecimiento '
    'de Olist. Los patrones de comportamiento del cliente en ese período pueden no representar '
    'fielmente el comportamiento en períodos más recientes o en otras etapas del ciclo de vida '
    'del negocio.')
bullet(doc, 'Se implementaron tres períodos de validación independientes (Val, BackTest, Live) para detectar degradación temporal.', 'Mitigación')
bullet(doc, 'El AUC mejora progresivamente (59.8% → 61% → 63%), indicando que el modelo no sobreaprende el período histórico.', 'Resultado')

heading2(doc, '3.2 Sesgo por clase desbalanceada')
body(doc,
    'El dataset presenta un fuerte desbalance: churn=0 (clientes que abandonan) representa '
    'solo el 1.2% del total, mientras que churn=1 (clientes activos) es el 98.8%. Sin '
    'corrección, el modelo tendería a ignorar la clase minoritaria.')
bullet(doc, 'Se aplicó SMOTE (oversampling) + RandomUnderSampler para balancear las clases durante el entrenamiento.', 'Mitigación')
bullet(doc, 'La evaluación se realizó con AUC-ROC, métrica robusta ante desbalance.', 'Métrica')

heading2(doc, '3.3 Sesgo de variables proxy')
body(doc,
    'Variables como installments_max (uso de cuotas) o freight_total_sum (costo de flete) '
    'pueden actuar como proxies de condición socioeconómica del cliente. El uso de estas '
    'variables para decisiones de retención podría resultar en tratamiento diferenciado '
    'basado indirectamente en el nivel de ingresos.')
bullet(doc, 'Monitorear que las tasas de predicción de churn no difieran significativamente entre regiones geográficas del Brasil.', 'Recomendación')
bullet(doc, 'Evaluar impacto diferencial por segmento de cliente en futuras iteraciones.', 'Futuro')

heading2(doc, '3.4 Sesgo de data leakage corregido')
body(doc,
    'Durante el desarrollo se detectó y corrigió un data leakage en la variable lead_min: '
    'se utilizaba la fecha de entrega real (posterior al corte temporal T) para calcular '
    'el tiempo de entrega mínimo. Tras la corrección, la importancia de lead_min bajó de '
    '~20% a ~7%, reflejando su contribución real.')
info_box(doc,
    'La corrección del leakage es un ejemplo de práctica ética: el modelo no debe usar '
    'información del futuro para predecir el presente, ya que esto generaría predicciones '
    'irreproducibles en producción real.', 'FFF3CD')

# ══════════════════════════════════════════════════════════
# 4. PRIVACIDAD Y LGPD
# ══════════════════════════════════════════════════════════
heading1(doc, '4. Privacidad y Cumplimiento LGPD')

heading2(doc, '4.1 Marco legal aplicable')
body(doc,
    'El dataset corresponde a clientes de Olist en Brasil. La Lei Geral de Proteção de '
    'Dados Pessoais (LGPD, Lei nº 13.709/2018) regula el tratamiento de datos personales '
    'en Brasil, con principios equivalentes al GDPR europeo. Entró en vigor en agosto de '
    '2020 con sanciones desde agosto de 2021.')

heading2(doc, '4.2 Datos utilizados y su clasificación')

# Tabla de datos
t_data = doc.add_table(rows=6, cols=3)
t_data.style = 'Table Grid'
headers_data = ['Variable / Grupo', 'Tipo de dato', 'Clasificación LGPD']
for j, h in enumerate(headers_data):
    cell = t_data.rows[0].cells[j]
    set_cell_bg(cell, '2E86AB')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)

filas_datos = [
    ('customer_unique_id',    'Identificador anonimizado',       'Dato personal seudonimizado'),
    ('Historial de órdenes',  'Fechas, montos, categorías',      'Dato personal — comportamiento'),
    ('review_max / buenas',   'Calificaciones de productos',     'Dato personal — opinión'),
    ('lead_min, freight_*',   'Tiempos y costos de entrega',     'Dato operacional — no sensible'),
    ('customer_zip_prefix',   'Prefijo de código postal',        'Dato personal — ubicación aproximada'),
]
for i, (v, t, c) in enumerate(filas_datos):
    row = t_data.rows[i+1]
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F0F4F8')
    for j, txt in enumerate([v, t, c]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10)
doc.add_paragraph()

heading2(doc, '4.3 Principios LGPD aplicados')
bullet(doc, 'El dataset de Olist es público (Kaggle). customer_unique_id es un hash que no permite reidentificación directa.', 'Anonimización')
bullet(doc, 'Las variables seleccionadas son de comportamiento de compra, no datos sensibles (salud, religión, biometría).', 'Minimización de datos')
bullet(doc, 'El modelo predice comportamiento futuro basado exclusivamente en historial de compras del propio cliente.', 'Finalidad')
bullet(doc, 'El dataset se usa únicamente para el desarrollo del modelo académico, no para decisiones comerciales reales.', 'Uso académico')

heading2(doc, '4.4 Consideraciones para producción real')
body(doc,
    'Si este modelo se desplegara en un entorno productivo real, sería necesario:')
bullet(doc, 'Obtener base legal explícita (consentimiento o legítimo interés) para el tratamiento predictivo.')
bullet(doc, 'Implementar el derecho de explicación: el cliente afectado por una decisión automatizada tiene derecho a conocer los motivos.')
bullet(doc, 'Registrar el tratamiento de datos en el Registro de Atividades de Tratamento (RAT).')
bullet(doc, 'Designar un Encarregado de Proteção de Dados (DPO) si el volumen de tratamiento lo requiere.')
bullet(doc, 'Realizar una Avaliação de Impacto à Proteção de Dados (AIPD) antes del despliegue.')

# ══════════════════════════════════════════════════════════
# 5. USO RESPONSABLE
# ══════════════════════════════════════════════════════════
heading1(doc, '5. Uso Responsable del Modelo')

heading2(doc, '5.1 Usos permitidos y recomendados')
bullet(doc, 'Priorizar clientes para campañas de retención proactiva (descuentos, contacto personalizado).')
bullet(doc, 'Segmentar la base de clientes por nivel de riesgo para optimizar el presupuesto de marketing.')
bullet(doc, 'Identificar factores operacionales que aumentan el churn (ej: tiempos de entrega) para mejorar el servicio.')
bullet(doc, 'Monitoreo continuo de la salud de la base de clientes.')

heading2(doc, '5.2 Usos NO recomendados')
bullet(doc, 'Negar servicios o beneficios a clientes clasificados como "churn alto" — el modelo predice probabilidad, no certeza.')
bullet(doc, 'Usar las predicciones como criterio único de decisión sin revisión humana.')
bullet(doc, 'Aplicar el modelo a segmentos de clientes muy diferentes al perfil de entrenamiento (ej: clientes B2B vs B2C).')
bullet(doc, 'Compartir predicciones individuales con terceros sin base legal.')

info_box(doc,
    'Principio clave: El modelo es una herramienta de apoyo a la decisión, no un sustituto '
    'del juicio humano. Un score alto de churn debe iniciar una conversación con el cliente, '
    'no una acción automatizada irreversible.', 'FDECEA')

# ══════════════════════════════════════════════════════════
# 6. GOBERNANZA Y MONITOREO
# ══════════════════════════════════════════════════════════
heading1(doc, '6. Gobernanza y Plan de Monitoreo')

heading2(doc, '6.1 Indicadores de monitoreo en producción')

t_mon = doc.add_table(rows=5, cols=3)
t_mon.style = 'Table Grid'
headers_mon = ['Indicador', 'Umbral de alerta', 'Acción recomendada']
for j, h in enumerate(headers_mon):
    cell = t_mon.rows[0].cells[j]
    set_cell_bg(cell, '1E3A5F')
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)

filas_mon = [
    ('AUC en producción',          'Caída > 5 puntos vs Live (63%)',  'Reentrenamiento inmediato'),
    ('PSI de features',            'PSI > 0.25 en cualquier variable', 'Revisar distribución + reentrenar'),
    ('Tasa de churn observada',    'Desviación > 2% vs histórico 1.2%', 'Análisis de causa raíz'),
    ('Lift Decil 1',               'Lift < 1.5x en período reciente',  'Revisión del modelo'),
]
for i, (ind, umbral, accion) in enumerate(filas_mon):
    row = t_mon.rows[i+1]
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F0F4F8')
    for j, txt in enumerate([ind, umbral, accion]):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10)
doc.add_paragraph()

heading2(doc, '6.2 Ciclo de reentrenamiento')
bullet(doc, 'Frecuencia recomendada: mensual, con los datos del mes más reciente.')
bullet(doc, 'Criterio de reentrenamiento obligatorio: PSI > 0.25 o caída de AUC > 5 puntos.')
bullet(doc, 'Evaluación en BackTest y Live antes de desplegar cada nueva versión.')
bullet(doc, 'Versionado del modelo: guardar pickle con fecha y métricas de evaluación.')

heading2(doc, '6.3 Responsabilidades')
bullet(doc, 'Científico de datos: mantenimiento del modelo, monitoreo de métricas, reentrenamiento.', 'Data Scientist')
bullet(doc, 'Equipo de negocio: definición de acciones de retención, validación de resultados en campo.', 'Negocio')
bullet(doc, 'DPO (si aplica): supervisión del cumplimiento LGPD, auditorías de uso del modelo.', 'DPO')
bullet(doc, 'Alta dirección: aprobación del uso del modelo para decisiones de alto impacto.', 'Dirección')

# ══════════════════════════════════════════════════════════
# 7. LIMITACIONES
# ══════════════════════════════════════════════════════════
heading1(doc, '7. Limitaciones del Modelo')
bullet(doc, 'AUC de 63% indica capacidad predictiva moderada. El modelo no predice con certeza, solo estima probabilidades.', 'Capacidad predictiva')
bullet(doc, 'Datos de 2016-2018 de una plataforma brasileña en crecimiento. Puede no generalizar a otros mercados o períodos.', 'Representatividad')
bullet(doc, 'El modelo no captura eventos externos (crisis económica, competencia nueva, cambios de política) que impacten el churn.', 'Variables externas')
bullet(doc, 'La tasa de churn del 1.2% es muy baja — pequeños errores de clasificación tienen gran impacto en métricas de precisión.', 'Desbalance extremo')
bullet(doc, 'El CLV y el costo de campaña usados en el análisis económico son estimaciones de mercado, no valores reales de Olist.', 'Impacto económico')

# ══════════════════════════════════════════════════════════
# 8. CONCLUSIONES
# ══════════════════════════════════════════════════════════
heading1(doc, '8. Conclusiones y Recomendaciones')
body(doc,
    'El modelo de churn desarrollado cumple con los principios éticos fundamentales para '
    'sistemas de machine learning aplicados a datos de clientes:')
bullet(doc, 'Transparente: utiliza Logistic Regression con coeficientes interpretables y análisis SHAP.')
bullet(doc, 'Honesto: evaluación en tres períodos independientes sin contaminación metodológica.')
bullet(doc, 'Robusto: corrige data leakage, controla overfitting y valida estabilidad temporal.')
bullet(doc, 'Accionable: Lift 2.24x permite campañas de retención 2x más eficientes que la selección aleatoria.')

doc.add_paragraph()
body(doc,
    'Para un despliegue responsable en producción, se recomienda implementar el plan de '
    'monitoreo descrito en la Sección 6, obtener las bases legales LGPD correspondientes '
    'y mantener siempre la supervisión humana sobre las decisiones de retención de clientes.')

info_box(doc,
    'El modelo es un punto de partida, no un oráculo. Su valor real está en orientar '
    'recursos hacia los clientes con mayor riesgo, no en reemplazar el criterio humano '
    'ni en automatizar decisiones que afecten derechos de los clientes.', 'E8F4FD')

# ─── FOOTER ───────────────────────────────────────────────
section = doc.sections[0]
footer  = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = fp.add_run('Informe Ético y de Gobernanza · Modelo de Churn Olist · Grupo 7 · Sprint 4')
r_foot.font.size  = Pt(9)
r_foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ─── GUARDAR ──────────────────────────────────────────────
output_path = 'informe_etico_gobernanza.docx'
doc.save(output_path)
print(f"✓ Documento guardado: {output_path}")
